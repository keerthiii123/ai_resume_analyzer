import os
import re
import io

import fitz
import streamlit as st

from PIL import Image
from dotenv import load_dotenv
from google import genai

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.colors import HexColor

from docx import Document
from docx.shared import Pt


# =====================================================
# GEMINI CONFIG
# =====================================================

load_dotenv()

try:
    api_key = st.secrets["GEMINI_API_KEY"]
except Exception:
    api_key = os.getenv("GEMINI_API_KEY")

client = genai.Client(api_key=api_key) if api_key else None

GEMINI_MODEL = "gemini-3.6-flash"


# =====================================================
# ROLE BASED SKILLS
# =====================================================

ROLE_SKILLS = {

    "AI Engineer": [
        "python",
        "langchain",
        "langgraph",
        "crewai",
        "rag",
        "chromadb",
        "fastapi",
        "huggingface",
        "embeddings",
        "vector database",
        "llm",
        "generative ai",
        "groq",
        "git",
        "github",
        "streamlit"
    ],

    "Python Developer": [
        "python",
        "django",
        "flask",
        "fastapi",
        "sql",
        "mysql",
        "postgresql",
        "rest api",
        "git",
        "oop"
    ],

    "Full Stack Developer": [
        "html",
        "css",
        "javascript",
        "react",
        "django",
        "python",
        "mysql",
        "node",
        "rest api",
        "git"
    ],

    "Data Analyst": [
        "python",
        "sql",
        "excel",
        "power bi",
        "tableau",
        "pandas",
        "numpy",
        "matplotlib",
        "seaborn"
    ],

    "Software Engineer": [
        "python",
        "java",
        "sql",
        "data structures",
        "algorithms",
        "git",
        "oop"
    ],

    "Accountant": [
        "tally",
        "gst",
        "sap",
        "excel",
        "accounting",
        "bookkeeping",
        "accounts payable",
        "accounts receivable"
    ],

    "HR / Recruiter": [
        "recruitment",
        "sourcing",
        "linkedin",
        "naukri",
        "screening",
        "excel"
    ],

    "Customer Support": [
        "customer support",
        "crm",
        "communication",
        "email support",
        "chat support"
    ]
}


# =====================================================
# SKILL ALIASES
# =====================================================

ALIASES = {

    "vector database": [
        "chromadb",
        "chroma db",
        "pinecone",
        "faiss",
        "weaviate",
        "vector db"
    ],

    "huggingface": [
        "hugging face",
        "sentence-transformers"
    ],

    "rag": [
        "retrieval augmented generation",
        "retrieval-augmented generation"
    ],

    "llm": [
        "large language model",
        "large language models"
    ],

    "github": [
        "github.com"
    ],

    "generative ai": [
        "gen ai",
        "genai",
        "generative artificial intelligence"
    ],

    "fastapi": [
        "fast api"
    ],

    "react": [
        "react.js",
        "react js"
    ],

    "javascript": [
        "java script"
    ],

    "postgresql": [
        "postgres",
        "postgre sql"
    ],

    "mysql": [
        "my sql"
    ],

    "power bi": [
        "powerbi",
        "power-bi"
    ],

    "rest api": [
        "restful api"
    ],

    "langchain": [
        "lang chain"
    ],

    "langgraph": [
        "lang graph"
    ],

    "chromadb": [
        "chroma db",
        "chroma"
    ]
}


# =====================================================
# TEXT NORMALIZATION
# =====================================================

def normalize_text(text):

    if not text:
        return ""

    text = text.lower()

    text = text.replace("–", "-")
    text = text.replace("—", "-")
    text = text.replace("•", " ")

    text = re.sub(r"\s+", " ", text)

    return text.strip()


# =====================================================
# PDF TEXT EXTRACTION
# =====================================================

def extract_text(uploaded_file):

    uploaded_file.seek(0)

    pdf = fitz.open(
        stream=uploaded_file.read(),
        filetype="pdf"
    )

    text = ""

    for page in pdf:
        text += page.get_text("text") + "\n"

    pdf.close()

    uploaded_file.seek(0)

    return text.strip()


# =====================================================
# RESUME PREVIEW
# =====================================================

def resume_preview(uploaded_file):

    uploaded_file.seek(0)

    pdf = fitz.open(
        stream=uploaded_file.read(),
        filetype="pdf"
    )

    if len(pdf) == 0:
        pdf.close()
        uploaded_file.seek(0)
        return None

    page = pdf.load_page(0)

    pix = page.get_pixmap(
        matrix=fitz.Matrix(2, 2)
    )

    image = Image.open(
        io.BytesIO(
            pix.tobytes("png")
        )
    )

    pdf.close()

    uploaded_file.seek(0)

    return image


# =====================================================
# SKILL DETECTION
# =====================================================

def skill_exists(text, skill):

    text = normalize_text(text)

    keywords = [
        skill
    ] + ALIASES.get(skill, [])

    for keyword in keywords:

        keyword = normalize_text(keyword)

        if not keyword:
            continue

        if len(keyword) <= 3:

            pattern = rf"\b{re.escape(keyword)}\b"

            if re.search(pattern, text):
                return True

        else:

            if keyword in text:
                return True

    return False


# =====================================================
# EXTRACT JD KEYWORDS
# =====================================================

def extract_jd_keywords(job_description):

    if not job_description.strip():
        return []

    jd_text = normalize_text(job_description)

    all_skills = sorted(
        set(
            skill
            for skills in ROLE_SKILLS.values()
            for skill in skills
        )
    )

    matched_jd_skills = []

    for skill in all_skills:

        if skill_exists(
            jd_text,
            skill
        ):

            matched_jd_skills.append(skill)

    return matched_jd_skills


# =====================================================
# JOB DESCRIPTION MATCH
# =====================================================

def job_description_match(
    resume_text,
    job_description
):

    if not job_description.strip():

        return 0, [], []

    resume_text = normalize_text(
        resume_text
    )

    jd_keywords = extract_jd_keywords(
        job_description
    )

    matched = []
    missing = []

    for keyword in jd_keywords:

        if skill_exists(
            resume_text,
            keyword
        ):

            matched.append(keyword)

        else:

            missing.append(keyword)

    total_keywords = (
        len(matched)
        + len(missing)
    )

    if total_keywords == 0:

        return 0, [], []

    percentage = round(
        len(matched)
        / total_keywords
        * 100
    )

    return (
        percentage,
        matched,
        missing
    )


# =====================================================
# ATS SCORE - 100
# =====================================================

def calculate_ats_score(
    resume_text,
    role,
    job_description=""
):

    text = normalize_text(
        resume_text
    )

    # =================================================
    # 1. ROLE SKILLS - 30
    # =================================================

    role_skills = ROLE_SKILLS.get(
        role,
        []
    )

    found_skills = []
    missing_skills = []

    for skill in role_skills:

        if skill_exists(
            text,
            skill
        ):

            found_skills.append(skill)

        else:

            missing_skills.append(skill)

    skills_score = round(
        len(found_skills)
        / max(1, len(role_skills))
        * 30
    )

    # =================================================
    # 2. JOB DESCRIPTION - 25
    # =================================================

    jd_percent, matched_keywords, missing_keywords = (
        job_description_match(
            resume_text,
            job_description
        )
    )

    jd_score = round(
        jd_percent
        * 25
        / 100
    )

    # =================================================
    # 3. EXPERIENCE - 15
    # =================================================

    experience_keywords = [
        "experience",
        "work experience",
        "professional experience",
        "employment",
        "internship"
    ]

    has_experience = any(
        keyword in text
        for keyword in experience_keywords
    )

    experience_score = (
        15
        if has_experience
        else 0
    )

    # =================================================
    # 4. PROJECTS - 10
    # =================================================

    project_keywords = [
        "project",
        "projects",
        "portfolio"
    ]

    has_projects = any(
        keyword in text
        for keyword in project_keywords
    )

    project_score = (
        10
        if has_projects
        else 0
    )

    # =================================================
    # 5. EDUCATION - 5
    # =================================================

    education_keywords = [
        "b.e",
        "b.e.",
        "be ",
        "btech",
        "b.tech",
        "bachelor",
        "master",
        "m.e",
        "m.e.",
        "mtech",
        "engineering",
        "university",
        "college",
        "degree"
    ]

    has_education = any(
        keyword in text
        for keyword in education_keywords
    )

    education_score = (
        5
        if has_education
        else 0
    )

    # =================================================
    # 6. RESUME SECTIONS - 5
    # =================================================

    sections = [
        "summary",
        "objective",
        "skills",
        "experience",
        "education",
        "projects"
    ]

    section_count = sum(
        1
        for section in sections
        if section in text
    )

    sections_score = round(
        section_count
        / len(sections)
        * 5
    )

    # =================================================
    # 7. CONTACT INFORMATION - 5
    # =================================================

    email_found = bool(
        re.search(
            r"[A-Za-z0-9._%+-]+@"
            r"[A-Za-z0-9.-]+\."
            r"[A-Za-z]{2,}",
            text
        )
    )

    phone_found = bool(
        re.search(
            r"\b\d{10}\b",
            text
        )
    )

    linkedin_found = (
        "linkedin" in text
    )

    contact_items = sum(
        [
            email_found,
            phone_found,
            linkedin_found
        ]
    )

    contact_score = round(
        contact_items
        / 3
        * 5
    )

    # =================================================
    # 8. ATS FORMATTING - 5
    # =================================================

    formatting_score = 0

    if len(text) > 500:

        formatting_score += 2

    if len(text.split()) >= 150:

        formatting_score += 1

    heading_keywords = [
        "skills",
        "experience",
        "education"
    ]

    heading_count = sum(
        1
        for heading in heading_keywords
        if heading in text
    )

    if heading_count >= 2:

        formatting_score += 2

    formatting_score = min(
        formatting_score,
        5
    )

    # =================================================
    # TOTAL SCORE
    # =================================================

    total = (
        skills_score
        + jd_score
        + experience_score
        + project_score
        + education_score
        + sections_score
        + contact_score
        + formatting_score
    )

    total = min(
        100,
        total
    )

    # =================================================
    # BREAKDOWN
    # =================================================

    breakdown = {

        "skills": skills_score,

        "job_description": jd_score,

        "education": education_score,

        "projects": project_score,

        "experience": experience_score,

        "sections": sections_score,

        "contact": contact_score,

        "formatting": formatting_score
    }

    return (
        total,
        found_skills,
        missing_skills,
        missing_skills[:10],
        breakdown
    )


# =====================================================
# GEMINI GENERATOR
# =====================================================

def gemini_generate(prompt):

    if client is None:

        return (
            "⚠️ GEMINI_API_KEY is not configured."
        )

    try:

        response = client.models.generate_content(
            model=GEMINI_MODEL,
            contents=prompt
        )

        if response and response.text:

            return response.text

        return (
            "⚠️ Empty response from Gemini."
        )

    except Exception as e:

        error = str(e)
        error_lower = error.lower()

        if (
            "api_key_invalid"
            in error_lower
            or "invalid api key"
            in error_lower
            or "401"
            in error_lower
            or "unauthenticated"
            in error_lower
        ):

            return (
                "⚠️ Invalid Gemini API key."
            )

        if "429" in error_lower:

            return (
                "⚠️ Gemini quota/rate limit exceeded."
            )

        if (
            "404"
            in error_lower
            or "not found"
            in error_lower
        ):

            return (
                f"⚠️ Gemini model "
                f"'{GEMINI_MODEL}' "
                "was not found."
            )

        return (
            f"⚠️ Gemini Error: {error}"
        )


# =====================================================
# AI RESUME ANALYSIS
# =====================================================

def analyze_resume(
    resume_text,
    role,
    job_description=""
):

    resume_text = resume_text[:12000]

    prompt = f"""
You are an expert ATS Resume Reviewer.

Target Role:
{role}

Job Description:
{
    job_description
    if job_description
    else "Not provided"
}

Analyze the resume carefully.

Return ONLY these sections:

# Professional Summary

# Strengths

# Weaknesses

# Suggested Improvements

# Five Interview Questions

Rules:

- Use only information available in the resume.
- Do not invent experience.
- Do not invent skills.
- Do not invent companies.
- Give practical ATS-focused suggestions.
- Consider the target role.
- If a job description is provided, consider relevant keywords.

Resume:
{resume_text}
"""

    return gemini_generate(
        prompt
    )


# =====================================================
# AI RESUME REWRITE
# =====================================================

def rewrite_resume(
    resume_text,
    role,
    job_description=""
):

    resume_text = resume_text[:12000]

    prompt = f"""
You are a professional ATS Resume Writer.

Rewrite the following resume professionally.

Target Role:
{role}

Job Description:
{
    job_description
    if job_description
    else "Not provided"
}

Rules:

- Use ONLY genuine information.
- Do NOT invent companies.
- Do NOT invent projects.
- Do NOT invent experience.
- Do NOT invent skills.
- Do NOT change dates.
- Do NOT create fake achievements.
- Improve grammar.
- Improve clarity.
- Use professional wording.
- Make it ATS-friendly.
- Use relevant keywords only when supported by the original resume.

Use these sections:

# Professional Summary

# Technical Skills

# Experience

# Projects

# Education

Resume:
{resume_text}
"""

    return gemini_generate(
        prompt
    )


# =====================================================
# PDF REPORT
# =====================================================

def generate_pdf_report(
    filename,
    role,
    score,
    found_skills,
    missing_skills,
    breakdown,
    ai_summary=""
):

    styles = getSampleStyleSheet()

    title = styles["Heading1"]

    title.alignment = TA_CENTER

    title.textColor = HexColor(
        "#2563EB"
    )

    doc = SimpleDocTemplate(
        filename
    )

    elements = []

    elements.append(
        Paragraph(
            "AI Resume ATS Report",
            title
        )
    )

    elements.append(
        Spacer(1, 20)
    )

    elements.append(
        Paragraph(
            f"<b>Target Role:</b> {role}",
            styles["BodyText"]
        )
    )

    elements.append(
        Paragraph(
            f"<b>ATS Score:</b> {score}/100",
            styles["BodyText"]
        )
    )

    elements.append(
        Spacer(1, 15)
    )

    elements.append(
        Paragraph(
            "Score Breakdown",
            styles["Heading2"]
        )
    )

    for key, value in breakdown.items():

        elements.append(
            Paragraph(
                f"{key.replace('_', ' ').title()}: {value}",
                styles["BodyText"]
            )
        )

    elements.append(
        Spacer(1, 15)
    )

    elements.append(
        Paragraph(
            "Skills Found",
            styles["Heading2"]
        )
    )

    elements.append(
        Paragraph(
            ", ".join(found_skills)
            if found_skills
            else "None",
            styles["BodyText"]
        )
    )

    elements.append(
        Spacer(1, 15)
    )

    elements.append(
        Paragraph(
            "Missing Skills",
            styles["Heading2"]
        )
    )

    elements.append(
        Paragraph(
            ", ".join(missing_skills)
            if missing_skills
            else "None",
            styles["BodyText"]
        )
    )

    if ai_summary:

        elements.append(
            Spacer(1, 15)
        )

        elements.append(
            Paragraph(
                "AI Summary",
                styles["Heading2"]
            )
        )

        safe_summary = (
            ai_summary
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )

        elements.append(
            Paragraph(
                safe_summary,
                styles["BodyText"]
            )
        )

    doc.build(
        elements
    )

    return filename


# =====================================================
# DOCX RESUME GENERATOR
# =====================================================

def generate_resume_docx(
    filename,
    rewritten_resume,
    role
):

    doc = Document()

    heading = doc.add_heading(
        "AI Rewritten Resume",
        1
    )

    if heading.runs:

        heading.runs[0].font.size = Pt(22)

    doc.add_paragraph(
        f"Target Role: {role}"
    )

    doc.add_paragraph()

    for line in rewritten_resume.splitlines():

        line = line.strip()

        if not line:
            continue

        if line.startswith("# "):

            doc.add_heading(
                line[2:].strip(),
                level=1
            )

        elif line.startswith("## "):

            doc.add_heading(
                line[3:].strip(),
                level=2
            )

        elif line.startswith("- "):

            doc.add_paragraph(
                line[2:].strip(),
                style="List Bullet"
            )

        elif re.match(
            r"^\d+\.\s+",
            line
        ):

            content = re.sub(
                r"^\d+\.\s+",
                "",
                line
            )

            doc.add_paragraph(
                content,
                style="List Number"
            )

        else:

            doc.add_paragraph(
                line
            )

    doc.save(
        filename
    )

    return filename

